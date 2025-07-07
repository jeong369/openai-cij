import random
import json
import os
from docx import Document
from uuid import uuid4

# 저장 위치
save_dir = "streamlit_data/ia_word_documents_50"
os.makedirs(save_dir, exist_ok=True)

# 샘플 과제명
task_titles = [
    "이메일 인증 절차 간소화", "결제 API 응답속도 개선", "모바일 UI 접근성 향상",
    "유저 세션 타임아웃 연장", "회원탈퇴 안내 메시지 강화", "보안로그 정밀 수집 기능",
    "다국어 지원 레이어 개선", "장애 알림 메시지 커스터마이징", "신규 결제수단 적용", "계정 잠금 정책 변경"
]

# 파트 정보
parts = [
    {"name": "인증관리팀", "role": "개발"},
    {"name": "UI/UX팀", "role": "테스트"},
    {"name": "결제시스템팀", "role": "개발"},
    {"name": "고객정보팀", "role": "운영"},
    {"name": "운영모니터링팀", "role": "운영"},
]

# 샘플 20개 생성
samples = []
for i in range(20):
    task_id = str(uuid4())[:8]
    title = random.choice(task_titles)
    requirement = f"{title} 관련하여 고객 요구가 반복적으로 발생하고 있어 개선이 요구됨."
    selected_parts = random.sample(parts, k=2)

    # Word 문서 생성
    doc = Document()
    doc.add_paragraph(f"[IA 문서] 과제번호: {task_id}")
    doc.add_paragraph(f"과제명: {title}")
    doc.add_paragraph(f"요구사항: {requirement}")
    for idx, p in enumerate(selected_parts, 1):
        doc.add_paragraph(f"[연관 파트{idx}]")
        doc.add_paragraph(f"- 파트 이름: {p['name']}")
        doc.add_paragraph(f"- 구분: {p['role']}")
        doc.add_paragraph(f"- 기능: {title} 관련 {'기능 수정' if p['role'] == '개발' else '테스트 수행'}")
        doc.add_paragraph(f"- 내용: {p['name']}에서 {title} 관련 작업을 수행함")
    
    filename = f"{task_id}_{title[:15]}.docx".replace(" ", "_")
    filepath = os.path.join(save_dir, filename)
    doc.save(filepath)

    # JSON 형태로도 반환
    sample = {
        "id": task_id,
        "project_title": title,
        "requirement": requirement,
        "parts": [
            {
                "part_name": p["name"],
                "role": p["role"],
                "feature": f"{title} 관련 {'기능 수정' if p['role'] == '개발' else '테스트 수행'}",
                "content": f"{p['name']}에서 {title} 관련 작업을 수행함"
            } for p in selected_parts
        ]
    }
    samples.append(sample)

# 저장된 Word 파일 경로 일부만 보여줌
samples[:3], f"{len(samples)}개 생성 완료"
