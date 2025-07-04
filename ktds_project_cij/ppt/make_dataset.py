from docx import Document
from docx.shared import Pt

# 문서 생성
doc = Document()
doc.add_heading('IA 문서 - Part D (온라인 화면 담당)', level=1)

sections = [
    ("1. 개요", """
Part D는 사용자 인터페이스(UI) 및 온라인 화면을 담당하는 파트입니다.
본 문서는 고객 요구사항에 따라 변경된 온라인 화면 구성 및 기능 명세를 정리한 IA 문서입니다.
"""),
    ("2. AS-IS", """
- 로그인 화면은 ID/PW 입력 후 로그인 버튼 제공
- 고객 정보 수정 화면은 한 페이지 내에 모든 필드가 나열되어 있음
- 결재 화면은 고정된 레이아웃(최대 3단계)으로 구성되어 있음
- 장애 발생 시 별도의 알림 없이 화면에 에러 메시지만 표시됨
"""),
    ("3. TO-BE", """
- 로그인 화면에 OTP 입력 필드 추가 및 UI 정비
- 고객 정보 수정 화면을 탭 구조로 변경하여 가독성 향상
- 결재 화면은 결재단계에 따라 동적 UI 적용(최대 5단계)
- 시스템 장애 발생 시 에러 메시지 외에 팝업 + 토스트 알림 표시
- 실시간 알림 UI: Slack 연동 상태 표시 영역 추가
"""),
    ("4. 변경 UI 시안", """
[로그인 화면]
- 기존: ID / PW 입력 후 로그인
- 변경: ID / PW / OTP 입력 → 로그인

[고객정보 수정 화면]
- 기존: 한 화면에 모든 항목 노출
- 변경: [기본정보][계약정보][변경이력] 탭 구성

[결재 화면]
- 기존: 고정된 결재 단계 UI
- 변경: 결재단계 수에 따라 동적으로 노드 UI 생성
"""),
    ("5. 영향 범위", """
- Part A (인증 관리): OTP 관련 UI 반영
- Part B (고객 정보 관리): 이력 탭 표시 연동
- Part C (결재/알림): 결재단계 동적 표현, 장애 알림 시각화

Part D는 위 세 파트의 기능 변경과 강하게 연관되어 있으며,
최종 사용자와 직접 상호작용하는 화면을 구성합니다.
"""),
    ("6. 변경 소스 및 구성 요소", """
- login.component.html: OTP 입력 필드 추가
- customer-edit.component.html: 탭 구성 HTML 적용
- approval.component.ts: 결재단계별 동적 UI 로직 추가
- error-popup.component.ts: 장애 발생 시 팝업 로직 구현
- slack-status.component.ts: Slack 상태 UI 구성 추가
"""),
    ("7. 테스트 및 검수 항목", """
- OTP 입력 오류 시 UX 흐름 정상 여부 확인
- 고객정보 수정 탭 전환 시 데이터 로딩 검증
- 결재단계 변경 시 UI 변경 반영 여부
- 장애 발생 시 알림 노출 형태 점검
"""),
    ("8. 추가 고려사항", """
- 접근성(A11y) 준수 여부 검토
- 모바일 화면 대응을 위한 반응형 디자인 적용
- 향후 알림 채널 확장(Google Chat, MS Teams 등) 고려
""")
]

# 내용 추가
for title, content in sections:
    doc.add_heading(title, level=2)
    for para in content.strip().split('\n'):
        p = doc.add_paragraph(para.strip())
        p.style.font.size = Pt(10.5)

# 저장
output_path = "/mnt/data/IA_DOC_PART_D.docx"
doc.save(output_path)

output_path
